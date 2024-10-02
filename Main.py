import streamlit as st
import os
import sqlite3
import pandas as pd
from Giris_ekrani import user_authentication  # Kullanıcı kimlik doğrulama
from streamlit_option_menu import option_menu
from datetime import datetime


# Veritabanı bağlantısı
def get_connection():
    conn = sqlite3.connect('database.db')
    return conn


# Sayfa yapılandırması
st.set_page_config(page_title="CV Oluşturucu", layout="wide")
CSV_FILES = {
    'personal_info': 'personal_info.csv',
    'education': 'education.csv',
    'projects': 'projects.csv',
    'experience': 'experience.csv',
    'certifications': 'certifications.csv',
    'activities': 'activities.csv',
    'skills': 'skills.csv',
    'accounts': 'accounts.csv'  # 'accounts' CSV dosyası
}



# Giriş durumu kontrolü
if 'login' not in st.session_state:
    st.session_state.login = False  # Oturum durumu

# Kullanıcı kimlik doğrulama
if not st.session_state.login:
    # Kullanıcı kimlik doğrulama işlemini başlat
    user_authenticated = user_authentication()

    # Kullanıcı giriş yaptı mı?
    if user_authenticated is not None and user_authenticated[0]:  # Eğer kullanıcı başarılı bir şekilde giriş yaptıysa
        st.session_state.login = True  # Oturum durumunu güncelle
        st.session_state.user_id = user_authenticated[1]  # Kullanıcı ID'sini sakla
        st.rerun()


# Eğer kullanıcı giriş yapmadıysa burada bir mesaj göster
if not st.session_state.login:
    st.error("Lütfen giriş yapın.")
    st.stop()  # İşlemi durdur, diğer bölümlere geçiş yapma


# Sol taraftaki menü
sidebar = st.sidebar
sidebar.title("CV Oluşturucu")

with sidebar:
    selected = option_menu("Ana Menü",
                           ["Kişisel Bilgiler", "Eğitim", "Projeler", "Deneyim",
                            "Sertifikalar", "Aktiviteler", "Beceriler", "CV Asistanı"],
                           icons=['person', 'book', 'projector', 'briefcase',
                                  'certificate', 'activity', 'tools', 'chat'],
                           menu_icon="cast", default_index=0)  # Varsayılan olarak "Kişisel Bilgiler" sekmesi açık

DATA_DIR =""
photo_folder_path=""



update_fields = {
        "PersonalInfo": ["info_id", "user_id", "name", "email", "phone"],
        "Photos": [ "photo_id", "user_id" , "photo"],  # Eğer bu tabloda güncellenebilir başka alan yoksa sadece "photo" yeterlidir
        "Education": ["edu_id", "user_id", "university", "faculty", "department", "start_year", "end_year"],
        "Projects": ["project_id", "user_id", "title", "description", "technologies", "show_in_cv"],
        "Experience": ["exp_id", "user_id", "company", "position", "start_date", "end_date", "description", "show_in_cv"],
        "Certifications": ["cert_id", "user_id", "name", "institution", "date", "show_in_cv"],
        "Activities": ["activity_id", "user_id", "name", "role", "date_range", "description", "show_in_cv"],
        "Skills": ["skill_id", "user_id", "name", "proficiency", "show_in_cv"],
        "Accounts": ["account_id", "user_id", "name", "link","show_in_cv"],
    }


def find_record_id_by_values(table_name, user_id, updated_entry):
    conn = get_connection()
    cursor = conn.cursor()

    # updated_entry'deki alanlara göre WHERE cümlesini oluştur
    where_clause = " AND ".join([f"{key} = ?" for key in updated_entry.keys()])

    # Sorguyu oluştur: user_id ve updated_entry değerlerine göre ilgili tabloyu sorgular
    query = f"""
        SELECT {update_fields[table_name][0]}  -- Anahtar sütunu seçiyoruz
        FROM {table_name} 
        WHERE user_id = ? AND {where_clause}
    """

    # Sorguyu çalıştır ve sonuçları al
    cursor.execute(query, (user_id, *updated_entry.values()))
    result = cursor.fetchone()

    # Bağlantıyı kapat
    conn.close()

    # Eğer sonuç bulunduysa ID'yi döndür, bulunmadıysa None döndür
    if result:
        return result[0]
    else:
        return None






def save_info(table_name, info):
    conn = get_connection()
    cursor = conn.cursor()

    # Tabloya özgü sütunları belirle
    if table_name == "PersonalInfo":
        columns = "user_id, name, email, phone"
    elif table_name == "Photos":
        columns = "user_id, photo"
    elif table_name == "Education":
        columns = "user_id, university, faculty, department, start_year, end_year"
    elif table_name == "Projects":
        columns = "user_id, title, description, technologies, show_in_cv"
    elif table_name == "Experience":
        columns = "user_id, company, position, start_date, end_date, description, show_in_cv"
    elif table_name == "Certifications":
        columns = "user_id, name, institution, date, show_in_cv"
    elif table_name == "Activities":
        columns = "user_id, name, role, date_range, description, show_in_cv"
    elif table_name == "Skills":
        columns = "user_id, name, proficiency, show_in_cv"
    elif table_name == "Accounts":
        columns = "user_id, name, link ,show_in_cv"
    else:
        raise ValueError("Geçersiz tablo adı.")

    # Sütunları ayırma ve değerlerin yerini alacak placeholder'ları (soru işareti) oluşturma
    placeholders = ', '.join(['?' for _ in info])

    # Kayıt ekleme sorgusu
    query = f"""
        INSERT INTO {table_name} ({columns}) 
        VALUES ({placeholders})
    """

    # Sözlükteki değerleri alarak sorguya gönderme
    cursor.execute(query, tuple(info.values()))
    conn.commit()
    conn.close()











# Fonksiyon: Diğer bölümler için genel veri silme




# Oturum durumu için başlangıç verileri
if 'edit_section' not in st.session_state:
    st.session_state['edit_section'] = None
if 'edit_index' not in st.session_state:
    st.session_state['edit_index'] = None



# Kullanıcının profil fotoğrafını veritabanından çeker
def get_first_photo_from_directory(user_id):
    """Kullanıcının profil fotoğrafını veritabanından çeker."""
    conn = get_connection()
    cursor = conn.cursor()

    # Kullanıcının fotoğrafını al
    cursor.execute("SELECT photo FROM Photos WHERE user_id = ?", (user_id,))
    photo = cursor.fetchone()  # Sadece bir fotoğraf döneceği için fetchone kullanıyoruz
    conn.close()

    if photo and photo[0]:
        return photo[0]  # BLOB formatındaki fotoğraf
    return None  # Eğer fotoğraf yoksa None döner




from Giris_ekrani import fetch_data

import os
import numpy as np
from PIL import Image  # Pillow kütüphanesini içeri aktarıyoruz


import sqlite3


def delete_entry(table_name, first_column_value):
    """Belirtilen tablo ve first_column_value ile bir kaydı siler."""


    # Eğer first_column_value None ise hata vermek yerine geri dön
    if first_column_value is None:
        st.write("Silinecek kayıt bulunamadı, first_column_value None.")
        return

    conn = get_connection()  # Veritabanı bağlantısını aç
    cursor = conn.cursor()

    # Her tablonun kendine özgü ID sütununa göre silme sorgusunu oluştur
    id_column_mapping = {
        "PersonalInfo": "info_id",
        "Photos": "photo_id",
        "Education": "edu_id",
        "Projects": "project_id",
        "Experience": "exp_id",
        "Certifications": "cert_id",
        "Activities": "activity_id",
        "Skills": "skill_id",
        "Accounts": "account_id"
    }

    # ID sütunu belirle
    id_column = id_column_mapping.get(table_name)
    if not id_column:
        raise ValueError("Geçersiz tablo adı.")

    # Silme sorgusunu oluştur
    query = f"""
    DELETE FROM {table_name}
    WHERE {id_column} = ?
    """

    # Sorguyu çalıştır
    cursor.execute(query, (first_column_value,))  # Tuple olarak geçiyoruz

    # Değişiklikleri kaydet ve bağlantıyı kapat
    conn.commit()
    conn.close()

    print("Silme işlemi başarılı.")

# ID sütunlarının tabloya göre eşleşmesi
id_column_mapping = {
    "PersonalInfo": "info_id",
    "Photos": "photo_id",
    "Education": "edu_id",
    "Projects": "project_id",
    "Experience": "exp_id",
    "Certifications": "cert_id",
    "Activities": "activity_id",
    "Skills": "skill_id",
    "Accounts": "account_id"
}

def update_show_in_cv_status(table_name, first_column_value, show_in_cv):
    """
    Belirtilen tablodaki bir girdinin `show_in_cv` alanını günceller.

    Args:
        table_name (str): Güncellenecek tablo adı.
        entry_id (int): Güncellenecek girdinin benzersiz kimliği.
        show_in_cv (bool): CV'de gösterilip gösterilmeyeceği durumu (True/False).
    """
    # Tablonun id sütununu id_column_mapping'den al
    id_column_mapping = {
        "PersonalInfo": "info_id",
        "Photos": "photo_id",
        "Education": "edu_id",
        "Projects": "project_id",
        "Experience": "exp_id",
        "Certifications": "cert_id",
        "Activities": "activity_id",
        "Skills": "skill_id",
        "Accounts": "account_id"
    }

    id_column = id_column_mapping.get(table_name)

    if not id_column:
        st.error(f"{table_name} için geçerli bir ID sütunu bulunamadı.")
        return

    # SQL güncelleme sorgusu (tablo adı ve id sütununa göre)
    query = f"""
        UPDATE {table_name}
        SET show_in_cv = ?
        WHERE {id_column} = ?
    """

    # Veritabanına bağlanıp sorguyu çalıştır
    conn = get_connection()  # Veritabanı bağlantısını aç
    with conn:
        conn.execute(query, (show_in_cv, first_column_value))



def update_entry(table_name, first_column_value, updated_entry):
    """Belirtilen tablo ve first_column_value ile güncelleme yapar."""
    conn = get_connection()  # Veritabanı bağlantısını aç
    cursor = conn.cursor()

    # Tabloya göre sütunları ve güncellenecek değerleri belirle
    if table_name == "PersonalInfo":
        query = f"UPDATE {table_name} SET name = ?, email = ?, phone = ? WHERE info_id = ?"
        values = (updated_entry['name'], updated_entry['email'], updated_entry['phone'], first_column_value)
    elif table_name == "Photos":
        query = f"UPDATE {table_name} SET photo = ? WHERE photo_id = ?"
        values = (updated_entry['photo'], first_column_value)
    elif table_name == "Education":
        query = f"UPDATE {table_name} SET university = ?, faculty = ?, department = ?, start_year = ?, end_year = ? WHERE edu_id = ?"
        values = (updated_entry['university'], updated_entry['faculty'], updated_entry['department'], updated_entry['start_year'], updated_entry['end_year'], first_column_value)
    elif table_name == "Projects":
        query = f"UPDATE {table_name} SET title = ?, description = ?, technologies = ?, show_in_cv = ? WHERE project_id = ?"
        values = (updated_entry['title'], updated_entry['description'], updated_entry['technologies'], updated_entry['show_in_cv'], first_column_value)
    elif table_name == "Experience":
        query = f"UPDATE {table_name} SET company = ?, position = ?, start_date = ?, end_date = ?, description = ?, show_in_cv = ? WHERE exp_id = ?"
        values = (updated_entry['company'], updated_entry['position'], updated_entry['start_date'], updated_entry['end_date'], updated_entry['description'], updated_entry['show_in_cv'], first_column_value)
    elif table_name == "Certifications":
        query = f"UPDATE {table_name} SET name = ?, institution = ?, date = ?, show_in_cv = ? WHERE cert_id = ?"
        values = (updated_entry['name'], updated_entry['institution'], updated_entry['date'], updated_entry['show_in_cv'], first_column_value)
    elif table_name == "Activities":
        query = f"UPDATE {table_name} SET name = ?, role = ?, date_range = ?, description = ?, show_in_cv = ? WHERE activity_id = ?"
        values = (updated_entry['name'], updated_entry['role'], updated_entry['date_range'], updated_entry['description'], updated_entry['show_in_cv'], first_column_value)
    elif table_name == "Skills":
        query = f"UPDATE {table_name} SET name = ?, proficiency = ?, show_in_cv = ? WHERE skill_id = ?"
        values = (updated_entry['name'], updated_entry['proficiency'], updated_entry['show_in_cv'], first_column_value)
    elif table_name == "Accounts":
        query = f"UPDATE {table_name} SET name = ?, link = ?, show_in_cv = ? WHERE account_id = ?"
        values = (updated_entry['name'], updated_entry['link'], updated_entry['show_in_cv'], first_column_value)
    else:
        raise ValueError("Geçersiz tablo adı.")

    # Güncelleme sorgusunu çalıştır
    cursor.execute(query, values)

    # Değişiklikleri kaydet ve bağlantıyı kapat
    conn.commit()
    conn.close()

    print("Güncelleme başarılı.")




def fetch_first_column_value_by_index(table_name, idx):
    """Belirtilen tablodaki idx değerine göre ilk kolonun değerini getirir."""
    conn = get_connection()
    cursor = conn.cursor()

    # Satırı almak için sorgu oluştur
    query = f"SELECT * FROM {table_name} LIMIT 1 OFFSET ?"
    cursor.execute(query, (idx,))

    # Sonucu al
    row = cursor.fetchone()

    # Eğer satır varsa, ilk kolonu döndür, yoksa None döndür
    if row:
        first_column_value = row[0]  # İlk kolonu al
    else:
        first_column_value = None  # Satır yoksa None

    # Bağlantıyı kapat
    conn.close()

    return first_column_value


# Veritabanından fotoğrafı alma fonksiyonu (matris formatında)

def convert_image_to_matrix(uploaded_file):
    """Yüklenen resmi matris formatına dönüştürür."""
    img = Image.open(uploaded_file)
    img = img.resize((150, 150))  # Resmi 150x150 boyutuna yeniden boyutlandır
    img_matrix = np.array(img)
    return img_matrix


def save_or_update_photo_matrix_to_db(user_id, photo_matrix):
    conn = get_connection()  # Veritabanı bağlantısını aç
    cursor = conn.cursor()

    # Fotoğrafı byte formatına çeviriyoruz
    photo_bytes = photo_matrix.tobytes()

    cursor.execute("SELECT photo FROM Photos WHERE user_id = ?", (user_id,))
    result = cursor.fetchone()

    if result:
        cursor.execute("UPDATE Photos SET photo = ? WHERE user_id = ?", (photo_bytes, user_id))
        st.write("Mevcut fotoğraf güncellendi.")
    else:
        cursor.execute("INSERT INTO Photos (user_id, photo) VALUES (?, ?)", (user_id, photo_bytes))
        st.write("Yeni fotoğraf kaydedildi.")

    conn.commit()
    conn.close()


def get_photo_matrix_from_db(user_id):
    conn = get_connection()  # Veritabanı bağlantısını aç
    cursor = conn.cursor()

    cursor.execute("SELECT photo FROM Photos WHERE user_id = ?", (user_id,))
    result = cursor.fetchone()

    conn.close()

    if result:
        photo_bytes = result[0]
        photo_matrix = np.frombuffer(photo_bytes, dtype=np.uint8)
        return photo_matrix
    return None


def matrix_to_image(photo_matrix, width=150, height=150, mode='RGB'):
    """Matrisi tekrar resme dönüştürür."""
    if photo_matrix is None or photo_matrix.size != width * height * 3:
        st.error("Görüntü boyutu hatalı veya boş.")
        return None
    img = Image.fromarray(photo_matrix.reshape((height, width, 3)), mode=mode)
    return img


def personal_info_section():
    st.header("Kişisel Bilgiler")

    table_name = "PersonalInfo"
    user_id = st.session_state.user_id

    # Veritabanından kişisel bilgileri getir
    personal_info = fetch_data(table_name, user_id)

    # Kişisel bilgilerin varlığı kontrolü
    name = personal_info['name'].iloc[0] if not personal_info.empty else ''
    email = personal_info['email'].iloc[0] if not personal_info.empty else ''
    phone = personal_info['phone'].iloc[0] if not personal_info.empty else ''

    if name and email and phone:
        buton_kisisel_bilgiler = "Kişisel Bilgileri Güncelle"
    else:
        buton_kisisel_bilgiler = "Kişisel Bilgileri Kaydet"

    # Kişisel Bilgiler Giriş Alanları
    name_input = st.text_input("İsim", name, key='personal_info_name')
    email_input = st.text_input("E-posta", email, key='personal_info_email')
    phone_input = st.text_input("Telefon (Sıfır olmadan)", phone, key='personal_info_phone',
                                placeholder="xxxxxxxxxx", max_chars=10)

    # Kişisel Bilgileri Kaydet Butonu
    if st.button(buton_kisisel_bilgiler):
        if not name_input.strip():
            st.error("İsim alanı zorunludur.")
        else:
            # Telefonu sayısal değerlere dönüştür
            phone_digits = ''.join(filter(str.isdigit, phone_input))
            if phone_digits and (len(phone_digits) == 10) and not phone_digits.startswith('0'):
                formatted_phone = f"({phone_digits[:3]}) {phone_digits[3:]}"

                # Kaydet veya Güncelle
                if buton_kisisel_bilgiler == "Kişisel Bilgileri Kaydet":
                    info = {
                        'user_id': user_id,
                        'name': name_input.strip(),
                        'email': email_input.strip(),
                        'phone': formatted_phone
                    }
                    save_info(table_name, info)  # Veritabanına kaydet

                elif buton_kisisel_bilgiler == "Kişisel Bilgileri Güncelle":

                    info = {
                        'name': name_input.strip(),
                        'email': email_input.strip(),
                        'phone': formatted_phone
                    }
                    # İlk kolonun değerini al (PersonalInfo için user_id olarak düşünülüyor)
                    update_entry("PersonalInfo", 1, info)


                #st.rerun()
            else:
                st.error("Geçersiz telefon numarası.")



    # Fotoğraf Yükleme Bölümü
    st.subheader("Fotoğraf")
    user_id = st.session_state.get('user_id', 1)  # Örnek kullanıcı ID, uygun şekilde değiştirin
    photo_matrix = get_photo_matrix_from_db(user_id)
    photo = st.file_uploader("Fotoğraf Yükle", type=["jpg", "jpeg", "png"], key="photo_uploader")

    if photo_matrix is not None:
        st.write("Mevcut CV Fotoğrafınız :")
        img = matrix_to_image(photo_matrix, width=150, height=150, mode='RGB')
        st.image(img, width=150)
        buton_foto_isim = "Fotoğrafı Güncelle"
    else:
        buton_foto_isim = "Fotoğrafı Kaydet"

    if photo is not None:
        st.write("Yeni Eklenen CV Fotoğrafı :")
        st.session_state['photo'] = photo
        st.image(photo, width=150)

    if st.button(buton_foto_isim) and 'photo' in st.session_state:
        uploaded_file = st.session_state['photo']
        if uploaded_file is not None:
            st.write("Yüklenen fotoğraf:", uploaded_file.name)

            photo_matrix = convert_image_to_matrix(uploaded_file)
            save_or_update_photo_matrix_to_db(user_id, photo_matrix)
            st.success(f"Fotoğraf başarıyla {buton_foto_isim.lower()}.")
            st.rerun()


    # Hesaplar Ekleme Bölümü
    st.subheader("Hesaplarınız")
    account_name = st.text_input("Sosyal Medya Adı", key="account_name")
    account_link = st.text_input("Hesap Linki", key="account_link")
    show_in_cv = st.checkbox("CV'de Göster", value=True, key="show_in_cv")

    # Hesap Ekle Butonu
    if st.button("Hesap Ekle", key='add_account_button'):
        if not account_name.strip() or not account_link.strip():
            st.error("Sosyal medya adı ve hesap linki zorunludur.")
        else:
            entry = {
                'user_id': user_id,
                'name': account_name.strip(),
                'link': account_link.strip(),
                'show_in_cv': show_in_cv
            }
            save_info('Accounts', entry)  # Hesabı veritabanına ekle
            st.success("Hesap eklendi.")
            st.rerun()

    # Eklenen Hesapları Listele
    st.subheader("Eklenen Hesaplar")
    df_accounts = fetch_data('Accounts', user_id)  # Kullanıcıya özel hesapları al
    if not df_accounts.empty:
        for idx, acc in df_accounts.iterrows():
            st.markdown(f"**{acc['name']}:** {acc['link']}")
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("Düzenle", key=f"edit_acc_{idx}_{acc['name']}"):
                    st.session_state['edit_section'] = 'accounts'
                    st.session_state['edit_index'] = idx
                    st.rerun()
            with col2:
                if st.button("Sil", key=f"delete_acc_{idx}_{acc['name']}"):
                    table_name = "Accounts"
                    first_column_value = fetch_first_column_value_by_index(table_name, idx)
                    delete_entry(table_name, first_column_value)  # Hesabı veritabanından sil
                    st.success("Hesap silindi.")
                    st.rerun()
            with col3:
                # Benzersiz bir key oluşturmak için idx ve hesap adı kullanıyoruz
                show_in_cv = st.checkbox("CV'de Göster", value=acc['show_in_cv'], key=f"show_activity_cv_{idx}")
                # Eğer checkbox durumu değişmişse, bunu veritabanına kaydedebilirsiniz
                first_column_value = fetch_first_column_value_by_index("Accounts", idx)
                update_show_in_cv_status('Accounts', first_column_value,
                                         show_in_cv)  # Veritabanını güncelleme fonksiyonu


    # Hesap Düzenleme
    if st.session_state.get('edit_section') == 'accounts' and st.session_state.get('edit_index') is not None:
        idx = st.session_state['edit_index']
        acc = df_accounts.iloc[idx]  # Düzenlenecek hesabı getir
        st.write(f"**{acc['name']}** düzenleniyor...")
        account_name_edit = st.text_input("Sosyal Medya Adı", acc['name'], key=f"acc_name_edit_{idx}")
        account_link_edit = st.text_input("Hesap Linki", acc['link'], key=f"acc_link_edit_{idx}")
        show_in_cv_edit = st.checkbox("CV'de Göster", value=acc['show_in_cv'], key=f"show_in_cv_edit_{idx}")

        if st.button("Kaydet", key=f"save_acc_edit_{idx}"):
            if not account_name_edit.strip() or not account_link_edit.strip():
                st.error("Sosyal medya adı ve hesap linki zorunludur.")
            else:
                updated_entry = {
                    'name': account_name_edit.strip(),
                    'link': account_link_edit.strip(),
                    'show_in_cv': show_in_cv_edit
                }
                table_name = "Accounts"
                # İlk kolonun değerini al
                first_column_value = fetch_first_column_value_by_index(table_name, idx)

                # Güncellemeyi gerçekleştir
                if first_column_value is not None:
                    update_entry(table_name, first_column_value, updated_entry)
                else:
                    print(f"Idx {idx} için güncellenebilecek bir satır bulunamadı.")

                st.success("Hesap güncellendi.")
                st.session_state['edit_section'] = None
                st.session_state['edit_index'] = None
                st.rerun()


        if st.button("İptal", key=f"cancel_acc_edit_{idx}"):
            st.session_state['edit_section'] = None
            st.session_state['edit_index'] = None
            st.rerun()




def education_section():
    st.header("Eğitim")

    table_name = "Education"
    user_id = st.session_state.user_id

    # Kayıtlı eğitim bilgilerini yükle
    df_education = fetch_data(table_name, user_id)

    if not df_education.empty:
        education = df_education.iloc[0].to_dict()
    else:
        education = {}

    # Kullanıcıdan eğitim bilgilerini alma
    university = st.text_input("Üniversite" )
    faculty = st.text_input("Fakülte" )
    department = st.text_input("Bölüm" )
    start_year = st.text_input("Başlangıç Yılı" )

    # Devam eden eğitim durumu
    ongoing = st.checkbox("Devam Ediyor",
                          key="education_ongoing")
    if not ongoing:
        end_year = st.text_input("Bitiş Yılı" )
    else:
        end_year = 'Devam Ediyor'

    # Kaydet butonu
    if st.button("Kaydet"):
        # Hata kontrolleri
        if not university.strip() or not faculty.strip() or not department.strip() or not start_year.strip():
            st.error("Tüm alanlar zorunludur.")
        elif not start_year.isdigit():
            st.error("Başlangıç Yılı sayısal bir değer olmalıdır.")
        elif not ongoing and (not end_year.strip() or not end_year.isdigit()):
            st.error("Bitiş Yılı sayısal bir değer olmalıdır.")
        else:
            # Eğitim bilgilerini veritabanına kaydet
            entry = {
                "user_id": user_id,
                'university': university.strip(),
                'faculty': faculty.strip(),
                'department': department.strip(),
                'start_year': start_year.strip(),
                'end_year': end_year.strip() if not ongoing else 'Devam Ediyor'
            }
            save_info("Education", entry)
            st.success("Eğitim bilgileri kaydedildi.")
            st.rerun()

    # Eklenen Eğitim Bilgilerini Listele
    st.subheader("Eklenen Eğitim Bilgileri")
    if not df_education.empty:
        for idx, row in df_education.iterrows():
            st.markdown(f"**Üniversite:** {row['university']} | **Fakülte:** {row['faculty']} | "
                        f"**Bölüm:** {row['department']} | **Başlangıç Yılı:** {row['start_year']} | "
                        f"**Bitiş Yılı:** {row['end_year']}")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Düzenle", key=f"edit_edu_{idx}"):
                    st.session_state['edit_section'] = 'education'
                    st.session_state['edit_index'] = idx
                    st.session_state['edit_entry'] = row.to_dict()  # Düzenlenecek kayıt bilgilerini sakla
                    st.rerun()  # Sayfayı yeniden yükle
            with col2:
                if st.button("Sil", key=f"delete_edu_{idx}"):
                    first_column_value = fetch_first_column_value_by_index(table_name, idx)
                    delete_entry(table_name, first_column_value)  # Eğitim bilgisini veritabanından sil
                    st.success("Eğitim bilgisi silindi.")
                    st.rerun()  # Sayfayı yeniden yükle

        # Eğitim Düzenleme
        if st.session_state.get('edit_section') == 'education' and st.session_state.get('edit_index') is not None:
            idx = st.session_state['edit_index']
            edu = st.session_state.get('edit_entry', df_education.iloc[idx])  # Düzenlenecek eğitim kaydını getir
            university_edit = st.text_input("Üniversite", edu['university'], key=f"uni_edit_{idx}")
            faculty_edit = st.text_input("Fakülte", edu['faculty'], key=f"fac_edit_{idx}")
            department_edit = st.text_input("Bölüm", edu['department'], key=f"dept_edit_{idx}")
            start_year_edit = st.text_input("Başlangıç Yılı", edu['start_year'], key=f"start_year_edit_{idx}")

            ongoing_edit = st.checkbox("Devam Ediyor", value=(edu['end_year'] == 'Devam Ediyor'),
                                       key=f"ongoing_edit_{idx}")
            if not ongoing_edit:
                end_year_edit = st.text_input("Bitiş Yılı", edu['end_year'], key=f"end_year_edit_{idx}")
            else:
                end_year_edit = 'Devam Ediyor'

            if st.button("Kaydet", key=f"save_edu_edit_{idx}"):
                if not university_edit.strip() or not faculty_edit.strip() or not department_edit.strip() or not start_year_edit.strip():
                    st.error("Tüm alanlar zorunludur.")
                elif not start_year_edit.isdigit():
                    st.error("Başlangıç Yılı sayısal bir değer olmalıdır.")
                elif not ongoing_edit and (not end_year_edit.strip() or not end_year_edit.isdigit()):
                    st.error("Bitiş Yılı sayısal bir değer olmalıdır.")
                else:
                    updated_entry = {
                        "user_id": user_id,
                        'university': university_edit.strip(),
                        'faculty': faculty_edit.strip(),
                        'department': department_edit.strip(),
                        'start_year': start_year_edit.strip(),
                        'end_year': end_year_edit.strip() if not ongoing_edit else 'Devam Ediyor'
                    }
                    # İlk kolonun değerini al
                    first_column_value = fetch_first_column_value_by_index(table_name, idx)
                    update_entry(table_name, first_column_value, updated_entry)
                    st.success("Eğitim bilgisi güncellendi.")
                    st.session_state['edit_section'] = None
                    st.session_state['edit_index'] = None
                    st.session_state['edit_entry'] = None  # Düzenleme bilgilerini temizle
                    st.rerun()  # Sayfayı yeniden yükle 1.kutucuğu değiştirdiğimde 2.kutucuğu değiştiremeden sayfa yenileniyor

        if st.button("İptal", key=f"cancel_edu_edit_{idx}"):
            st.session_state['edit_section'] = None
            st.session_state['edit_index'] = None
            st.rerun()







def projects_section():
    st.header("Projeler")

    table_name = "Projects"
    user_id = st.session_state.user_id

    # Kayıtlı projeleri yükle
    df_projects = fetch_data(table_name, user_id)

    if not df_projects.empty:
        project = df_projects.iloc[0].to_dict()
    else:
        project = {}

    # Yeni proje ekleme bölümünü oluştur
    project_title = st.text_input("Proje Başlığı (İsteğe Bağlı)", "")
    project_description = st.text_area("Proje Açıklaması", "")
    technologies = st.text_input("Kullanılan Teknolojiler (virgülle ayırın)", "")
    show_in_cv = st.checkbox("CV'de Göster", value=True)

    # Kaydet butonu
    if st.button("Kaydet"):
        if not project_description.strip():
            st.error("Proje açıklaması zorunludur.")
        else:
            # Proje bilgilerini veritabanına kaydet
            entry = {
                "user_id": user_id,
                'title': project_title.strip(),
                'description': project_description.strip(),
                'technologies': technologies.strip(),
                'show_in_cv': show_in_cv
            }
            save_info("Projects", entry)
            st.success("Proje bilgileri kaydedildi.")
            st.rerun()

    # Eklenen projeleri listele
    # Eklenen projeleri listele
    st.subheader("Eklenen Projeler")
    if not df_projects.empty:
        for idx, row in df_projects.iterrows():
            st.markdown(f"**Proje Başlığı:** {row['title']} | **Açıklama:** {row['description']} | "
                        f"**Kullanılan Teknolojiler:** {row['technologies']}")

            col1, col2, col3 = st.columns(3)

            # Düzenleme butonu
            with col1:
                if st.button("Düzenle", key=f"edit_proj_{idx}"):
                    st.session_state['edit_section'] = 'projects'
                    st.session_state['edit_index'] = idx

            # Silme butonu
            with col2:
                if st.button("Sil", key=f"delete_proj_{idx}"):
                    first_column_value = fetch_first_column_value_by_index("Projects", idx)
                    delete_entry("Projects", first_column_value)  # Projeyi veritabanından sil
                    st.success("Proje bilgisi silindi.")
                    st.rerun()  # Sayfayı yeniden yükle

            # CV'de göster checkbox
            with col3:
                # Benzersiz bir key oluşturmak için idx ve proje başlığı kullanılıyor
                show_in_cv = st.checkbox("CV'de Göster", value=row['show_in_cv'], key=f"show_proj_cv_{idx}")
                if show_in_cv != row['show_in_cv']:
                    # Eğer checkbox durumu değişmişse, bunu veritabanına kaydediyoruz
                    first_column_value = fetch_first_column_value_by_index("Projects", idx)
                    update_show_in_cv_status('Projects', first_column_value, show_in_cv)  # Veritabanını güncelle



        # Proje Düzenleme
        if st.session_state.get('edit_section') == 'projects' and st.session_state.get('edit_index') is not None:
            idx = st.session_state['edit_index']
            proj = df_projects.iloc[idx]  # Düzenlenecek proje kaydını getir
            project_title_edit = st.text_input("Proje Başlığı (İsteğe Bağlı)", proj['title'], key=f"title_edit_{idx}")
            project_description_edit = st.text_area("Proje Açıklaması", proj['description'], key=f"desc_edit_{idx}")
            technologies_edit = st.text_input("Kullanılan Teknolojiler", proj['technologies'], key=f"tech_edit_{idx}")
            show_in_cv_edit = st.checkbox("CV'de Göster", value=proj['show_in_cv'], key=f"show_edit_{idx}")

            if st.button("Kaydet", key=f"save_proj_edit_{idx}"):
                if not project_description_edit.strip():
                    st.error("Proje açıklaması zorunludur.")
                else:
                    updated_entry = {
                        'title': project_title_edit.strip(),
                        'description': project_description_edit.strip(),
                        'technologies': technologies_edit.strip(),
                        'show_in_cv': show_in_cv_edit
                    }

                    # İlk kolonun değerini al
                    first_column_value = fetch_first_column_value_by_index(table_name, idx)
                    update_entry("Projects", first_column_value, updated_entry)
                    st.success("Proje bilgisi güncellendi.")
                    st.session_state['edit_section'] = None  # Durumu sıfırla
                    st.session_state['edit_index'] = None  # Durumu sıfırla
                    st.rerun()  # Sayfayı yeniden yükle

            if st.button("İptal", key=f"cancel_proj_edit_{idx}"):
                st.session_state['edit_section'] = None  # Durumu sıfırla
                st.session_state['edit_index'] = None  # Durumu sıfırla
                st.rerun()  # Sayfayı yeniden yükle



def experience_section():
    st.header("Deneyim")

    table_name = "Experience"
    user_id = st.session_state.user_id

    # Kayıtlı deneyim bilgilerini yükle
    df_experience = fetch_data(table_name, user_id)

    if not df_experience.empty:
        experience = df_experience.iloc[0].to_dict()
    else:
        experience = {}

    # Kullanıcıdan deneyim bilgilerini alma
    company = st.text_input("Şirket Adı"  )
    position = st.text_input("Pozisyon" )
    start_date = st.date_input("Başlangıç Tarihi",
                                key="start_date")

    ongoing = st.checkbox("Devam Ediyor",  key="experience_ongoing")
    if ongoing:
        end_date = 'Devam Ediyor'
    else:
        end_date = st.date_input("Bitiş Tarihi",
                                  key="end_date").strftime("%d.%m.%Y")

    description = st.text_area("Deneyim Açıklaması" )
    show_in_cv = st.checkbox("CV'de Göster" ,value= True )

    # Kaydet butonu
    if st.button("Kaydet"):
        # Hata kontrolleri
        if not company.strip() or not position.strip() or not description.strip():
            st.error("Tüm alanlar zorunludur.")
        else:
            # Deneyim bilgilerini veritabanına kaydet
            entry = {
                "user_id": user_id,
                'company': company.strip(),
                'position': position.strip(),
                'start_date': start_date.strftime("%d.%m.%Y"),
                'end_date': end_date,
                'description': description.strip(),
                'show_in_cv': show_in_cv
            }
            save_info(table_name, entry)
            st.success("Deneyim bilgileri kaydedildi.")
            st.rerun()

    # Eklenen Deneyimleri Listele
    st.subheader("Eklenen Deneyimler")
    if not df_experience.empty:
        for idx, row in df_experience.iterrows():
            st.markdown(f"**Şirket:** {row['company']} | **Pozisyon:** {row['position']} | "
                        f"**Başlangıç Tarihi:** {row['start_date']} | **Bitiş Tarihi:** {row['end_date']}")
            st.markdown(f"**Açıklama:** {row['description']}")

            # Üç sütun oluşturuyoruz
            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("Düzenle", key=f"edit_exp_{idx}"):
                    st.session_state['edit_section'] = 'experience'
                    st.session_state['edit_index'] = idx
                    st.rerun()

            with col2:
                if st.button("Sil", key=f"delete_exp_{idx}"):
                    first_column_value = fetch_first_column_value_by_index("Experience", idx)
                    delete_entry("Experience", first_column_value)  # Deneyimi veritabanından sil
                    st.success("Deneyim bilgisi silindi.")
                    st.rerun()  # Sayfayı yeniden yükle

            with col3:
                # Benzersiz bir key oluşturmak için idx ve şirket adını kullanıyoruz
                show_in_cv = st.checkbox("CV'de Göster", value=row['show_in_cv'], key=f"show_exp_cv_{idx}")
                if show_in_cv != row['show_in_cv']:
                    # Eğer checkbox durumu değişmişse, bunu veritabanına kaydediyoruz
                    first_column_value = fetch_first_column_value_by_index("Experience", idx)
                    update_show_in_cv_status('Experience', first_column_value, show_in_cv)  # Veritabanını güncelle

    # Deneyim Düzenleme
    if st.session_state.get('edit_section') == 'experience' and st.session_state.get('edit_index') is not None:
        idx = st.session_state['edit_index']
        exp = df_experience.iloc[idx]  # Düzenlenecek deneyim kaydını getir

        # Kullanıcıdan güncellenmiş deneyim bilgilerini alma
        company_edit = st.text_input("Şirket Adı", exp['company'], key=f"company_edit_{idx}")
        position_edit = st.text_input("Pozisyon", exp['position'], key=f"position_edit_{idx}")
        start_date_edit = st.date_input("Başlangıç Tarihi",
                                         datetime.strptime(exp['start_date'], "%d.%m.%Y"), key=f"start_date_edit_{idx}")

        ongoing_edit = st.checkbox("Devam Ediyor", value=(exp['end_date'] == 'Devam Ediyor'), key=f"ongoing_edit_{idx}")
        if ongoing_edit:
            end_date_edit = 'Devam Ediyor'
        else:
            end_date_edit = st.date_input("Bitiş Tarihi",
                                           datetime.strptime(exp['end_date'], "%d.%m.%Y"), key=f"end_date_edit_{idx}").strftime("%d.%m.%Y")

        description_edit = st.text_area("Deneyim Açıklaması", exp['description'], key=f"description_edit_{idx}")
        show_in_cv_edit = st.checkbox("CV'de Göster", value=exp['show_in_cv'], key=f"show_cv_edit_{idx}")


        if st.button("Kaydet", key=f"save_exp_edit_{idx}"):
            if not company_edit.strip() or not position_edit.strip() or not description_edit.strip():
                st.error("Tüm alanlar zorunludur.")
            else:
                updated_entry = {
                    'company': company_edit.strip(),
                    'position': position_edit.strip(),
                    'start_date': start_date_edit.strftime("%d.%m.%Y"),
                    'end_date': end_date_edit,
                    'description': description_edit.strip(),
                    'show_in_cv': show_in_cv_edit
                }
                # İlk kolonun değerini al
                first_column_value = fetch_first_column_value_by_index(table_name, idx)
                st.write("değerrr : ",first_column_value)
                update_entry("Experience", first_column_value, updated_entry)
                st.success("Deneyim bilgisi güncellendi.")
                st.session_state['edit_section'] = None
                st.session_state['edit_index'] = None
                st.rerun()

        if st.button("İptal", key=f"cancel_exp_edit_{idx}"):
            st.session_state['edit_section'] = None
            st.session_state['edit_index'] = None
            st.rerun()


def certifications_section():
    st.header("Sertifikalar")

    table_name = "Certifications"
    user_id = st.session_state.user_id

    # Kayıtlı sertifika bilgilerini yükle
    df_certifications = fetch_data(table_name, user_id)

    if not df_certifications.empty:
        certification = df_certifications.iloc[0].to_dict()
    else:
        certification = {}

    # Kullanıcıdan sertifika bilgilerini alma
    certificate_name = st.text_input("Sertifika Adı" )
    institution = st.text_input("Kurum")



    year = st.text_input("Yıl")
    month = st.text_input("Ay (İsteğe Bağlı)")
    day = st.text_input("Gün (İsteğe Bağlı)" )
    show_in_cv = st.checkbox("CV'de Göster", value=certification.get('show_in_cv', True))

    # Kaydet butonu
    if st.button("Kaydet"):
        # Hata kontrolleri
        if not certificate_name.strip() or not institution.strip() or not year.strip():
            st.error("Sertifika adı, kurum ve yıl zorunludur.")
        else:
            if not year.isdigit():
                st.error("Yıl sayısal bir değer olmalıdır.")
            else:
                # Tarih formatı
                date = f"{day}.{month}.{year}" if day and month else (f"{month}.{year}" if month else year)
                cert_entry = {
                    'user_id': user_id,
                    'name': certificate_name.strip(),
                    'institution': institution.strip(),
                    'date': date,
                    'show_in_cv': show_in_cv
                }
                save_info(table_name, cert_entry)  # Sertifika bilgilerini kaydet
                st.success("Sertifika bilgileri kaydedildi.")
                st.rerun()

    # Eklenen Sertifikaları Listele
    st.subheader("Eklenen Sertifikalar")
    if not df_certifications.empty:
        for idx, cert in df_certifications.iterrows():
            st.markdown(
                f"**Sertifika Adı:** {cert['name']} | **Kurum:** {cert['institution']} | **Tarih:** {cert['date']}")

            # Üç sütun oluşturuyoruz
            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("Düzenle", key=f"edit_cert_{idx}"):
                    st.session_state['edit_section'] = 'Certifications'
                    st.session_state['edit_index'] = idx
                    st.rerun()

            with col2:
                if st.button("Sil", key=f"delete_cert_{idx}"):
                    first_column_value = fetch_first_column_value_by_index("Certifications", idx)
                    delete_entry("Certifications", first_column_value)  # Sertifikayı veritabanından sil
                    st.success("Sertifika bilgisi silindi.")
                    st.rerun()  # Sayfayı yeniden yükle

            with col3:
                # Benzersiz bir key oluşturmak için idx ve sertifika adını kullanıyoruz
                show_in_cv = st.checkbox("CV'de Göster", value=cert['show_in_cv'], key=f"show_cert_cv_{idx}")
                if show_in_cv != cert['show_in_cv']:
                    # Eğer checkbox durumu değişmişse, bunu veritabanına kaydediyoruz
                    first_column_value = fetch_first_column_value_by_index("Certifications", idx)
                    update_show_in_cv_status('Certifications', first_column_value, show_in_cv)  # Veritabanını güncelle

    # Sertifika Düzenleme
    if st.session_state.get('edit_section') == 'Certifications' and st.session_state.get('edit_index') is not None:
        idx = st.session_state['edit_index']
        cert = df_certifications.iloc[idx]

        # Kullanıcıdan güncellenmiş sertifika bilgilerini alma
        certificate_name_edit = st.text_input("Sertifika Adı", cert['name'], key=f"cert_name_edit_{idx}")
        institution_edit = st.text_input("Kurum", cert['institution'], key=f"institution_edit_{idx}")

        date_str_edit = cert.get('date', '')
        if date_str_edit:
            date_parts = date_str_edit.split('.')
            day_edit = date_parts[0] if len(date_parts) > 2 else ''
            month_edit = date_parts[1] if len(date_parts) > 1 else ''
            year_edit = date_parts[2] if len(date_parts) > 0 else ''
        else:
            day_edit, month_edit, year_edit = '', '', ''

        year_edit = st.text_input("Yıl", year_edit, key=f"year_edit_{idx}")
        month_edit = st.text_input("Ay (İsteğe Bağlı)", month_edit, key=f"month_edit_{idx}")
        day_edit = st.text_input("Gün (İsteğe Bağlı)", day_edit, key=f"day_edit_{idx}")
        show_in_cv_edit = st.checkbox("CV'de Göster", value=cert['show_in_cv'], key=f"show_cert_edit_{idx}")

        if st.button("Kaydet", key=f"save_cert_edit_{idx}"):
            if not certificate_name_edit.strip() or not institution_edit.strip() or not year_edit.strip():
                st.error("Sertifika adı, kurum ve yıl zorunludur.")
            else:
                if not year_edit.isdigit():
                    st.error("Yıl sayısal bir değer olmalıdır.")
                else:
                    date_edit = f"{day_edit}.{month_edit}.{year_edit}" if day_edit and month_edit else (f"{month_edit}.{year_edit}" if month_edit else year_edit)
                    updated_entry = {
                        'name': certificate_name_edit.strip(),
                        'institution': institution_edit.strip(),
                        'date': date_edit,
                        'show_in_cv': show_in_cv_edit
                    }
                    first_column_value = fetch_first_column_value_by_index(table_name, idx)
                    update_entry(table_name, first_column_value, updated_entry)
                    st.success("Sertifika bilgisi güncellendi.")
                    st.session_state['edit_section'] = None
                    st.session_state['edit_index'] = None
                    st.rerun()

        if st.button("İptal", key=f"cancel_cert_edit_{idx}"):
            st.session_state['edit_section'] = None
            st.session_state['edit_index'] = None
            st.rerun()


def activities_section():
    st.header("Aktiviteler")

    table_name = "Activities"
    user_id = st.session_state.user_id

    # Mevcut aktiviteleri yükle
    df_activities = fetch_data(table_name, user_id)

    # Yeni aktivite ekleme
    community_name = st.text_input("Topluluk Adı", key="community_name")
    role = st.text_input("Pozisyon", key="activity_role")
    start_date = st.date_input("Başlangıç Tarihi", value=datetime.today(), key="activity_start_date")
    ongoing = st.checkbox("Devam Ediyor", value=False, key="activity_ongoing")

    if not ongoing:
        end_date = st.date_input("Bitiş Tarihi", value=datetime.today(), key="activity_end_date")
        date_range = f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}"
    else:
        date_range = f"{start_date.strftime('%d.%m.%Y')} - Devam Ediyor"

    activity_description = st.text_area("Açıklama", key="activity_description")
    show_in_cv = st.checkbox("CV'de Göster", value=True, key="activity_show_cv")

    if st.button("Ekle"):
        if not community_name.strip() or not role.strip():
            st.error("Topluluk adı ve pozisyon zorunludur.")
        else:
            activity_entry = {
                'user_id': user_id,
                'name': community_name.strip(),
                'role': role.strip(),
                'date_range': date_range,
                'description': activity_description.strip(),
                'show_in_cv': show_in_cv
            }
            save_info('Activities', activity_entry)
            st.success("Aktivite eklendi.")
            st.rerun()

    # Eklenen Aktiviteler
    st.subheader("Eklenen Aktiviteler")
    if not df_activities.empty:
        for idx, act in df_activities.iterrows():
            st.markdown(f"**{act['name']} - {act['role']}**")
            st.markdown(f"   • {act['description']}")

            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("Düzenle", key=f"edit_activity_{idx}_{act['name']}"):
                    st.session_state['edit_section'] = 'Activities'
                    st.session_state['edit_index'] = idx
                    st.rerun()

            with col2:
                if st.button("Sil", key=f"delete_activity_{idx}_{act['name']}"):
                    first_column_value = fetch_first_column_value_by_index('Activities', idx)
                    delete_entry('Activities', first_column_value)  # Aktiviteleri veritabanından sil
                    st.success("Aktivite bilgisi silindi.")  # Hata düzeltildi
                    st.rerun()  # Sayfayı yeniden yükle

            with col3:
                show = st.checkbox("CV'de Göster", value=act['show_in_cv'], key=f"show_activity_{idx}_{act['name']}")
                if show != act['show_in_cv']:
                    # Eğer checkbox durumu değişmişse, bunu veritabanına kaydediyoruz
                    first_column_value = fetch_first_column_value_by_index('Activities', idx)
                    update_show_in_cv_status('Activities', first_column_value, show)  # Veritabanını güncelle

    # Eğer düzenleme modundaysa
    if st.session_state.get('edit_section') == 'Activities' and st.session_state.get('edit_index') is not None:
        idx = st.session_state['edit_index']
        act = df_activities.iloc[idx]
        st.write(f"**{act['name']} - {act['role']}** düzenleniyor...")

        # Mevcut bilgileri düzenle
        community_name = st.text_input("Topluluk Adı", act['name'], key=f"act_name_edit_{idx}")
        role = st.text_input("Pozisyon", act['role'], key=f"role_edit_{idx}")

        # Başlangıç ve Bitiş Tarihi Ayarları
        date_parts = act['date_range'].split(' - ')
        start_date = datetime.strptime(date_parts[0], "%d.%m.%Y")
        ongoing = date_parts[1] == 'Devam Ediyor'
        start_date = st.date_input("Başlangıç Tarihi", start_date, key=f"start_date_edit_{idx}")

        ongoing = st.checkbox("Devam Ediyor", value=ongoing, key=f"ongoing_edit_{idx}")
        if not ongoing:
            if len(date_parts) > 1 and date_parts[1] != 'Devam Ediyor':
                end_date = datetime.strptime(date_parts[1], "%d.%m.%Y")
            else:
                end_date = datetime.today()
            end_date = st.date_input("Bitiş Tarihi", end_date, key=f"end_date_edit_{idx}")
            end_date_str = end_date.strftime("%d.%m.%Y")
        else:
            end_date_str = 'Devam Ediyor'

        activity_description = st.text_area("Açıklama", act['description'], key=f"description_edit_{idx}")
        show_in_cv = st.checkbox("CV'de Göster", value=act['show_in_cv'], key=f"show_act_edit_{idx}")

        if st.button("Kaydet", key=f"save_act_edit_{idx}"):
            if not community_name.strip() or not role.strip():
                st.error("Topluluk adı ve pozisyon zorunludur.")
            else:
                date_range = f"{start_date.strftime('%d.%m.%Y')} - {'Devam Ediyor' if ongoing else end_date_str}"
                updated_entry = {
                    'name': community_name.strip(),
                    'role': role.strip(),
                    'date_range': date_range,
                    'description': activity_description.strip(),
                    'show_in_cv': show_in_cv
                }

                first_column_value = fetch_first_column_value_by_index(table_name, idx)
                update_entry('Activities', first_column_value, updated_entry)

                st.success("Aktivite güncellendi.")
                st.session_state['edit_section'] = None
                st.session_state['edit_index'] = None
                st.rerun()

        if st.button("İptal", key=f"cancel_act_edit_{idx}"):
            st.session_state['edit_section'] = None
            st.session_state['edit_index'] = None
            st.rerun()


def skills_section():
    st.header("Beceriler")

    table_name = "Skills"
    user_id = st.session_state.user_id

    # Mevcut becerileri yükle
    df_skills = fetch_data(table_name, user_id)

    # Yeni beceri ekleme
    skill_name = st.text_input("Beceri Adı", key="skill_name")
    proficiency = st.text_input("Yeterlilik Seviyesi", key="proficiency")
    show_in_cv = st.checkbox("CV'de Göster", value=True, key="show_skill_cv")

    if st.button("Ekle"):
        if not skill_name.strip():
            st.error("Beceri adı zorunludur.")
        else:
            skill_entry = {
                'user_id': user_id,
                'name': skill_name.strip(),
                'proficiency': proficiency.strip(),
                'show_in_cv': show_in_cv
            }
            save_info('Skills', skill_entry)
            st.success("Beceri eklendi.")
            st.rerun()

    # Eklenen Beceriler
    st.subheader("Eklenen Beceriler")
    if not df_skills.empty:
        for idx, skill in df_skills.iterrows():
            st.markdown(f"**{skill['name']}** - {skill['proficiency']}")

            col1, col2, col3 = st.columns(3)

            with col1:
                if st.button("Düzenle", key=f"edit_skill_{idx}_{skill['name']}"):
                    st.session_state['edit_section'] = 'Skills'
                    st.session_state['edit_index'] = idx
                    st.rerun()

            with col2:
                if st.button("Sil", key=f"delete_skill_{idx}_{skill['name']}"):
                    first_column_value = fetch_first_column_value_by_index('Skills', idx)
                    delete_entry('Skills', first_column_value)  # Beceriyi veritabanından sil
                    st.success("Beceri silindi.")  # Doğru mesaj verildi
                    st.session_state['edit_section'] = None  # Düzenleme bölümünü sıfırla
                    st.session_state['edit_index'] = None  # Seçilen indeksi sıfırla
                    st.rerun()  # Sayfayı yeniden yükle

            with col3:
                # CV'de göster checkbox'ı için kod eklenebilir
                show_in_cv = st.checkbox("CV'de Göster", value=skill['show_in_cv'], key=f"show_skill_{idx}")
                if show_in_cv != skill['show_in_cv']:
                    # Durum değiştiğinde veritabanını güncelle
                    first_column_value = fetch_first_column_value_by_index('Skills', idx)
                    update_show_in_cv_status('Skills', first_column_value, show_in_cv)  # Veritabanını güncelle



    # Eğer düzenleme modundaysa
    if st.session_state.get('edit_section') == 'Skills' and st.session_state.get('edit_index') is not None:
        idx = st.session_state['edit_index']
        skill = df_skills.iloc[idx]

        # Mevcut bilgileri düzenle
        skill_name = st.text_input("Beceri Adı", skill['name'], key=f"skill_name_edit_{idx}")
        proficiency = st.text_input("Yeterlilik Seviyesi", skill['proficiency'], key=f"proficiency_edit_{idx}")
        show_in_cv = st.checkbox("CV'de Göster", value=skill['show_in_cv'], key=f"show_skill_edit_{idx}")

        if st.button("Kaydet", key=f"save_skill_edit_{idx}"):
            if not skill_name.strip():
                st.error("Beceri adı zorunludur.")
            else:
                updated_entryy = {
                    'name': skill_name.strip(),
                    'proficiency': proficiency.strip(),
                    'show_in_cv': show_in_cv
                }

                first_column_value = fetch_first_column_value_by_index(table_name, idx)
                update_entry('Skills', first_column_value, updated_entryy)
                st.success("Beceri güncellendi.")
                st.session_state['edit_section'] = None
                st.session_state['edit_index'] = None
                st.rerun()

        if st.button("İptal", key=f"cancel_skill_edit_{idx}"):
            st.session_state['edit_section'] = None
            st.session_state['edit_index'] = None
            st.rerun()



from Chat_bot import run_cv_bot_app

section_functions = {
    "Kişisel Bilgiler": personal_info_section,
    "Eğitim": education_section,
    "Projeler": projects_section,
    "Deneyim": experience_section,
    "Sertifikalar": certifications_section,
    "Aktiviteler": activities_section,
    "Beceriler": skills_section,
    "CV Asistanı": run_cv_bot_app
}

selected_function = section_functions.get(selected)
if selected_function:
    selected_function()  # Fonksiyonu çağır
else:
    st.error("Seçilen bölüm bulunamadı.")  # Hata durumu için

from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
import os
import tempfile
from PIL import Image  # Resim işleme için Pillow kütüphanesi



def save_cv_as_word(markdown_content, filename):
    doc = Document()
    # Markdown içeriğini parçalara ayır
    sections = markdown_content.split('\n\n')
    for section in sections:
        # Başlıkları kontrol et
        if section.startswith('# '):  # Ana başlık
            doc.add_heading(section[2:], level=1)
        elif section.startswith('## '):  # Alt başlık
            doc.add_heading(section[3:], level=2)
        else:
            doc.add_paragraph(section)
    doc.save(filename)


from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
import streamlit as st
import pandas as pd
import tempfile

from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
import streamlit as st
import pandas as pd
import tempfile

from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
import streamlit as st
import pandas as pd
import tempfile


import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
import pandas as pd
import streamlit as st

def create_word_document():
    doc = Document()

    # 1. Kişisel Bilgiler
    table_name = "PersonalInfo"
    user_id = st.session_state.user_id

    personal_info_df = fetch_data(table_name, user_id)  # Veritabanından kişisel bilgileri yükle
    personal_info = personal_info_df.iloc[0].to_dict() if not personal_info_df.empty else {}

    # Kişisel bilgiler ve fotoğraf aynı hizada olacak şekilde tablo oluştur
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    # Sol sütun: Kişisel bilgiler (isim, e-posta, telefon vb.)
    cell1 = table.cell(0, 0)
    cell1.width = Inches(4)

    # İsim
    if 'name' in personal_info and personal_info['name']:
        name_paragraph = cell1.add_paragraph(personal_info['name'].upper())
        name_paragraph.runs[0].bold = True  # İsim kalın
    else:
        name_paragraph = cell1.add_paragraph("Ad Bilgisi Eksik")
        name_paragraph.runs[0].bold = True  # Ad eksikse de kalın yap

    # Telefon ve e-posta bilgilerini tek bir paragrafta birleştir
    phone = personal_info.get('phone', 'Telefon bilgisi eksik.')
    email = personal_info.get('email', 'E-posta bilgisi eksik.')
    contact_info_paragraph = cell1.add_paragraph(f"Telefon: {phone}                                                                ")
    contact_info_paragraph.add_run(f"E-posta: {email}")

    # Hesaplar
    table_name = "Accounts"
    df_accounts = fetch_data(table_name, user_id)

    if not df_accounts.empty:
        for acc in df_accounts.itertuples(index=False):
            # Hesap bilgilerini her birini yeni bir satıra ekle
            contact_info_paragraph.add_run(f"\n{acc.name}: {acc.link}")  # Hesapları yeni satırda ekle
            contact_info_paragraph.space_after = Pt(0)  # Boşluğu azalt

    # Sağ sütun: Fotoğraf
    cell2 = table.cell(0, 1)
    cell2.width = Inches(2)

    # Fotoğraf veritabanından çekiliyor
    user_id = st.session_state.get('user_id', 1)  # Kullanıcı ID, uygun şekilde ayarlayın
    photo_matrix = get_photo_matrix_from_db(user_id)

    if photo_matrix is not None:
        # Matrisi görüntüye çevir ve Word belgesine ekle
        img = matrix_to_image(photo_matrix, width=150, height=150, mode='RGB')

        # Geçici dosyaya kaydet ve Word belgesine ekle
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            img.save(tmp_file.name)
            photo_path = tmp_file.name
            # Fotoğrafı hücreye ekle
            cell2.add_paragraph().add_run().add_picture(photo_path, width=Inches(1.5))
    else:
        cell2.add_paragraph("Fotoğraf mevcut değil.")

    # 3. Eğitim Bilgileri
    table_name = "Education"
    df_education = fetch_data(table_name, user_id)

    if not df_education.empty:
        education = df_education.iloc[0]
        doc.add_heading("EĞİTİM", level=2)  # Başlık kalın ve büyük
        doc.add_paragraph(
            f"{education.get('university', 'Üniversite bilgisi yok')}, "
            f"{education.get('faculty', 'Fakülte bilgisi yok')} – "
            f"{education.get('department', 'Bölüm bilgisi yok')}, "
            f"{education.get('start_year', 'Başlangıç yılı yok')} - "
            f"{education.get('end_year', 'Bitiş yılı yok')}"
        )

    # 4. Projeler
    table_name = "Projects"
    df_projects = fetch_data(table_name, user_id)

    if 'show_in_cv' not in df_projects.columns:
        df_projects['show_in_cv'] = False  # show_in_cv eksikse False olarak eklenir

    projects_to_show = df_projects[df_projects['show_in_cv'] == True]
    if not projects_to_show.empty:
        doc.add_heading("PROJELER", level=2)  # Başlık kalın ve büyük
        for proj in projects_to_show.itertuples(index=False):
            if pd.notnull(proj.title) and len(proj.title) >= 3:
                doc.add_paragraph(f"• {proj.title}")
                doc.add_paragraph(f"   {proj.description}")
            else:
                doc.add_paragraph(f"• {proj.description}")
            if pd.notnull(proj.technologies) and len(proj.technologies) >= 3:
                doc.add_paragraph(f"   Teknolojiler: {proj.technologies}")

    # 5. Deneyim
    table_name = "Experience"
    df_experience = fetch_data(table_name, user_id)

    if 'show_in_cv' not in df_experience.columns:
        df_experience['show_in_cv'] = False  # show_in_cv eksikse False olarak eklenir

    experiences_to_show = df_experience[df_experience['show_in_cv'] == True]
    if not experiences_to_show.empty:
        doc.add_heading("DENEYİM", level=2)  # Başlık kalın ve büyük
        for exp in experiences_to_show.itertuples(index=False):
            doc.add_paragraph(f"• {exp.company} – {exp.position} ({exp.start_date} - {exp.end_date})")
            doc.add_paragraph(f"   {exp.description}")

    # 6. Sertifikalar
    table_name = "Certifications"
    df_certifications = fetch_data(table_name, user_id)

    if 'show_in_cv' not in df_certifications.columns:
        df_certifications['show_in_cv'] = False  # show_in_cv eksikse False olarak eklenir

    certifications_to_show = df_certifications[df_certifications['show_in_cv'] == True]
    if not certifications_to_show.empty:
        doc.add_heading("SERTİFİKALAR", level=2)  # Başlık kalın ve büyük
        for cert in certifications_to_show.itertuples(index=False):
            doc.add_paragraph(f"{cert.name}, {cert.institution} – {cert.date}")

    # 7. Aktiviteler
    table_name = "Activities"
    df_activities = fetch_data(table_name, user_id)

    if 'show_in_cv' not in df_activities.columns:
        df_activities['show_in_cv'] = False  # show_in_cv eksikse False olarak eklenir

    activities_to_show = df_activities[df_activities['show_in_cv'] == True]
    if not activities_to_show.empty:
        doc.add_heading("AKTİVİTELER", level=2)  # Başlık kalın ve büyük
        for act in activities_to_show.itertuples(index=False):
            doc.add_paragraph(f"• {act.name} – {act.role} ({act.date_range})")
            doc.add_paragraph(f"   {act.description}")

    # 8. Beceriler
    table_name = "Skills"
    df_skills = fetch_data(table_name, user_id)

    if 'show_in_cv' not in df_skills.columns:
        df_skills['show_in_cv'] = False  # show_in_cv eksikse False olarak eklenir

    skills_to_show = df_skills[df_skills['show_in_cv'] == True]
    if not skills_to_show.empty:
        doc.add_heading("BECERİLER", level=2)  # Başlık kalın ve büyük
        for skill in skills_to_show.itertuples(index=False):
            doc.add_paragraph(f"• {skill.name}: {skill.proficiency}")

    # Başlıkların stilini ayarla
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 2':
            for run in paragraph.runs:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk

    return doc





import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO



# Streamlit application code
# Streamlit uygulama kodu
if selected != "CV Asistanı":
    with st.container():
        col1, col2 = st.columns([4, 1])

        with col1:
            st.title("CV Taslağı ")
            st.write("Not : bilgiler girildikçe taslak oluşmaktadır ")

        with col2:
            # Word belgesini oluştur
            doc = create_word_document()

            # Belgeyi bellekte tutmak için BytesIO kullanın
            doc_stream = BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)

            # Gizli indirme butonu
            st.download_button(
                label="CV'yi İndir",
                data=doc_stream,
                file_name="CV.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="hidden_download_button",
                help="CV'nizi indirmek için tıklayın.",
                disabled=False,
                use_container_width=True,
                type="secondary"
            )

        # Sütunlar oluştur
        col1, col2 = st.columns([3, 1])

        with col1:
            table_name = "PersonalInfo"
            user_id = st.session_state.user_id
            personal_info_df = fetch_data(table_name, user_id)

            if personal_info_df.empty:
                pass
            else:
                personal_info = personal_info_df.iloc[0].to_dict()

                st.header(personal_info.get('name', '').upper())
                if personal_info.get('phone'):
                    st.write(f"**Telefon:** {personal_info.get('phone', '')}")

                st.write(f"**E-posta:** {personal_info.get('email', '')}")

                # Hesaplar
                table_name = "Accounts"
                df_accounts = fetch_data(table_name, user_id)

                if df_accounts.empty:
                    pass
                else:
                    if 'show_in_cv' not in df_accounts.columns:
                        df_accounts['show_in_cv'] = True  # Varsayılan olarak tüm hesapları göster

                    filtered_accounts = df_accounts[df_accounts['show_in_cv'] == True]

                    if not filtered_accounts.empty:
                        for acc in filtered_accounts.itertuples(index=False):
                            st.write(f"**{acc.name}:** {acc.link}")
                    else:
                        st.write("Gösterilecek hesap bulunmamaktadır.")

            # Eğitim
            table_name = "Education"
            df_education = fetch_data(table_name, user_id)

            if df_education.empty:
                pass
            else:
                education = df_education.iloc[0]
                st.header("EĞİTİM")
                st.write(
                    f"**{education.get('university', '')}, {education.get('faculty', '')}** – "
                    f"{education.get('department', '')}, {education.get('start_year', '')} - "
                    f"{education.get('end_year', '')}"
                )

            # Projeler
            table_name = "Projects"
            df_projects = fetch_data(table_name, user_id)

            if df_projects.empty:
                pass
            else:
                if 'show_in_cv' not in df_projects.columns:
                    df_projects['show_in_cv'] = True  # Varsayılan olarak tüm projeleri göster

                projects_to_show = df_projects[df_projects['show_in_cv'] == True]
                if not projects_to_show.empty:
                    st.header("PROJELER")
                    for proj in projects_to_show.itertuples(index=False):
                        title_display = proj.title if pd.notna(proj.title) else " "
                        technologies_display = proj.technologies if pd.notna(proj.technologies) else " "
                        if title_display != "":
                            st.markdown(f"• **{title_display}**")
                            st.markdown(f"   {proj.description}")
                        if proj.description != "":
                            st.markdown(f"•   {proj.description}")

            # Deneyim
            table_name = "Experience"
            df_experience = fetch_data(table_name, user_id)

            if df_experience.empty:
                pass
                experiences_sorted = pd.DataFrame()  # Boş bir DataFrame kullanın
            else:
                if 'end_date' not in df_experience.columns:
                    st.warning("Deneyim verilerinde 'end_date' sütunu bulunamadı.")
                    experiences_sorted = pd.DataFrame()  # Boş bir DataFrame kullanın
                else:
                    if 'show_in_cv' not in df_experience.columns:
                        df_experience['show_in_cv'] = True  # Varsayılan olarak tüm deneyimleri göster

                    experiences_to_show = df_experience[df_experience['show_in_cv'] == True]

                    def parse_date_cv(date_str):
                        if date_str == 'Devam Ediyor':
                            return datetime.max
                        else:
                            return datetime.strptime(date_str, "%d.%m.%Y")

                    experiences_sorted = experiences_to_show.copy()
                    experiences_sorted['sort_date'] = experiences_sorted['end_date'].apply(parse_date_cv)
                    experiences_sorted.sort_values(by=['sort_date', 'start_date'], ascending=False, inplace=True)

            if not experiences_sorted.empty:
                st.header("DENEYİM")
                for exp in experiences_sorted.itertuples(index=False):
                    st.markdown(f"• **{exp.company} – {exp.position}** ({exp.start_date} - {exp.end_date})")
                    st.markdown(f"   {exp.description}")

            # Sertifikalar
            table_name = "Certifications"
            df_certifications = fetch_data(table_name, user_id)

            if df_certifications.empty:
                pass
            else:
                if 'show_in_cv' not in df_certifications.columns:
                    df_certifications['show_in_cv'] = True  # Varsayılan olarak tüm sertifikaları göster

                certifications_to_show = df_certifications[df_certifications['show_in_cv'] == True]

                if not certifications_to_show.empty:
                    certifications_to_show['date'] = certifications_to_show['date'].fillna('None')
                    valid_dates = certifications_to_show['date'].astype(str).str.split('.', expand=True)

                    # Eğer valid_dates 3 sütundan daha fazla içeriyorsa, fazla sütunları atın.
                    if valid_dates.shape[1] == 3:
                        valid_dates = valid_dates[[0, 1]]  # Yıl ve Ay'ı alıyoruz (0: yıl, 1: ay)
                    elif valid_dates.shape[1] < 2:
                        # Eğer tarih eksikse, varsayılan bir değer verin
                        valid_dates['year'] = certifications_to_show['date'].str.split('.', expand=True)[0]
                        valid_dates['month'] = 1  # Ay eksikse varsayılan olarak Ocak'ı ekleyin

                    # Ardından sütun isimlerini atayın
                    valid_dates.columns = ['year', 'month']

                    certifications_to_show = certifications_to_show[valid_dates['year'] != 'None']
                    certifications_to_show[['year', 'month']] = valid_dates[['year', 'month']].astype(int)
                    certifications_to_show = certifications_to_show.sort_values(by=['year', 'month'], ascending=False)

                    st.header("SERTİFİKALAR")
                    for cert in certifications_to_show.itertuples(index=False):
                        st.markdown(f"• **{cert.name}**, {cert.institution} – {cert.date}")

            # Aktiviteler
            table_name = "Activities"
            df_activities = fetch_data(table_name, user_id)

            if df_activities.empty:
                pass
            else:
                if 'show_in_cv' not in df_activities.columns:
                    df_activities['show_in_cv'] = True  # Varsayılan olarak tüm aktiviteleri göster

                activities_to_show = df_activities[df_activities['show_in_cv'] == True]
                if not activities_to_show.empty:
                    st.header("AKTİVİTELER")
                    for act in activities_to_show.itertuples(index=False):
                        st.markdown(f"• **{act.name} – {act.role}** ({act.date_range})")
                        st.markdown(f"   {act.description}")

            # Beceriler
            table_name = "Skills"
            df_skills = fetch_data(table_name, user_id)

            if df_skills.empty:
                pass
            else:
                if 'show_in_cv' not in df_skills.columns:
                    df_skills['show_in_cv'] = True  # Varsayılan olarak tüm becerileri göster

                skills_to_show = df_skills[df_skills['show_in_cv'] == True]
                if not skills_to_show.empty:
                    st.header("BECERİLER")
                    for skill in skills_to_show.itertuples(index=False):
                        st.markdown(f"• **{skill.name}:** {skill.proficiency}")

        with col2:
            if st.session_state.get('photo'):
                st.image(st.session_state['photo'], width=150)
            else:
                user_id = st.session_state.get('user_id', 1)  # Örnek kullanıcı ID, uygun şekilde değiştirin
                photo_matrix = get_photo_matrix_from_db(user_id)

                if photo_matrix is not None:
                    img = matrix_to_image(photo_matrix, width=150, height=150, mode='RGB')
                    st.image(img)
                else:
                    pass
